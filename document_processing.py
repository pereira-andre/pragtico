import os
import re
import unicodedata
from datetime import datetime, timezone
from typing import Dict, List

from docx import Document
from pypdf import PdfReader
from werkzeug.utils import secure_filename


ALLOWED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx", ".csv"}


def slugify(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value).encode("ascii", "ignore").decode("ascii")
    cleaned = re.sub(r"[^a-zA-Z0-9]+", "-", normalized.strip().lower()).strip("-")
    return cleaned or "documento"


def extension_for(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


def is_allowed_document(filename: str) -> bool:
    return extension_for(filename) in ALLOWED_EXTENSIONS


def is_text_editable(filename: str) -> bool:
    return extension_for(filename) in {".md", ".txt"}


def sanitize_upload_filename(filename: str) -> str:
    safe_name = secure_filename(filename or "")
    if not safe_name:
        return f"documento{extension_for(filename) or '.txt'}"
    return safe_name


def ensure_unique_filename(directory: str, filename: str) -> str:
    stem, suffix = os.path.splitext(filename)
    candidate = filename
    counter = 1
    while os.path.exists(os.path.join(directory, candidate)):
        candidate = f"{stem}-{counter}{suffix}"
        counter += 1
    return candidate


def read_text_file(path: str) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            with open(path, "r", encoding=encoding) as handle:
                return handle.read()
        except UnicodeDecodeError:
            continue
    with open(path, "r", encoding="utf-8", errors="ignore") as handle:
        return handle.read()


def read_pdf_file(path: str) -> str:
    reader = PdfReader(path)
    pages: List[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"[Página {index}]\n{text}")
    return "\n\n".join(pages)


def read_docx_file(path: str) -> str:
    doc = Document(path)
    parts: List[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def read_csv_file(path: str) -> str:
    rows: List[str] = []
    with open(path, "r", encoding="utf-8", errors="ignore") as handle:
        header = handle.readline().strip()
        if header:
            rows.append(f"Cabeçalho CSV: {header}")
        for index, line in enumerate(handle, start=1):
            clean = line.strip()
            if clean:
                rows.append(clean)
            if index >= 120:
                rows.append("CSV truncado para indexação inicial.")
                break
    return "\n".join(rows)


def extract_text_from_path(path: str) -> str:
    suffix = extension_for(path)
    if suffix in {".md", ".txt"}:
        return read_text_file(path)
    if suffix == ".pdf":
        return read_pdf_file(path)
    if suffix == ".docx":
        return read_docx_file(path)
    if suffix == ".csv":
        return read_csv_file(path)
    raise ValueError(f"Extensão não suportada: {suffix}")


def build_preview(text: str, limit: int = 220) -> str:
    collapsed = re.sub(r"\s+", " ", text).strip()
    if len(collapsed) <= limit:
        return collapsed
    return collapsed[: limit - 1].rstrip() + "…"


def iso_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def format_bytes(size: int) -> str:
    units = ["B", "KB", "MB", "GB"]
    value = float(size)
    for unit in units:
        if value < 1024 or unit == units[-1]:
            if unit == "B":
                return f"{int(value)} {unit}"
            return f"{value:.1f} {unit}"
        value /= 1024
    return f"{size} B"


def infer_document_type(filename: str) -> str:
    suffix = extension_for(filename)
    if suffix == ".pdf":
        return "PDF"
    if suffix == ".docx":
        return "Word"
    if suffix == ".md":
        return "Markdown"
    if suffix == ".txt":
        return "Texto"
    if suffix == ".csv":
        return "CSV"
    return suffix.lstrip(".").upper() or "Ficheiro"


def file_metadata(path: str) -> Dict:
    stat = os.stat(path)
    return {
        "size_bytes": stat.st_size,
        "updated_at": datetime.fromtimestamp(stat.st_mtime, tz=timezone.utc).isoformat(),
    }
